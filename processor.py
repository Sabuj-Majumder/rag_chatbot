import os
import pandas as pd
import pdfplumber
import docx
import pytesseract
import sqlite3
from PIL import Image
from io import BytesIO
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document



class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )

    def process_file(self, file_path: str, file_type: str, original_filename: str):
        text_content = ""
        
        try:
            # 1. Handle PDF
            if file_type == "application/pdf" or file_path.endswith(".pdf"):
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        # Extract text
                        text = page.extract_text() or ""
                        text_content += text + "\n"
                        
            
            # 2. Handle DOCX
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_path.endswith(".docx"):
                doc = docx.Document(file_path)
                text_content = "\n".join([para.text for para in doc.paragraphs])

            # 3. Handle Text
            elif file_type == "text/plain" or file_path.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    text_content = f.read()

            # 4. Handle Images (OCR)
            elif file_type.startswith("image/") or file_path.endswith((".jpg", ".png", ".jpeg")):
                image = Image.open(file_path)
                text_content = pytesseract.image_to_string(image)

            # 5. Handle CSV
            elif file_type == "text/csv" or file_path.endswith(".csv"):
                df = pd.read_csv(file_path)
                text_content = df.to_string(index=False)

            # 6. Handle SQLite 
            elif file_path.endswith(".db") or file_path.endswith(".sqlite"):
                conn = sqlite3.connect(file_path)
                cursor = conn.cursor()
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                tables = cursor.fetchall()
                for table in tables:
                    table_name = table[0]
                    text_content += f"\nTable: {table_name}\n"
                    df = pd.read_sql_query(f"SELECT * FROM {table_name}", conn)
                    text_content += df.to_string(index=False) + "\n"
                conn.close()

            else:
                return []

            # Chunking
            docs = [Document(page_content=text_content, metadata={"source": original_filename})]
            return self.text_splitter.split_documents(docs)

        except Exception as e:
            print(f"Error processing file: {e}")
            return []